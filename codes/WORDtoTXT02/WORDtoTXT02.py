import os
import re
from docx import Document
from collections import Counter

def extract_text_from_paragraphs(paragraphs, common_headers):
    text = []
    for para in paragraphs:
        clean_text = para.text.strip()
        if clean_text and clean_text not in common_headers:
            text.append(clean_text)
    return "\n".join(text)

def extract_text_from_table(table):
    table_text = []
    for row in table.rows:
        row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
        if row_text:
            table_text.append(" | ".join(row_text))
    return "\n".join(table_text)

def remove_extra_newlines(text):
    return re.sub(r'\n{3,}', '\n', text)

def find_common_headers(docx_input, threshold=3):
    para_texts = [para.text.strip() for para in docx_input.paragraphs if para.text.strip()]
    para_counter = Counter(para_texts)
    common_headers = {para for para, count in para_counter.items() if count >= threshold}
    return common_headers

def docx_to_txt(input_path,base_path):
    docx_input = Document(input_path)
    filename = os.path.splitext(os.path.basename(input_path))[0]

    table_folder_path = f"{filename}_table"
    os.makedirs(os.path.join(base_path,table_folder_path), exist_ok=True)

    common_headers = find_common_headers(docx_input)

    paragraph_text = extract_text_from_paragraphs(docx_input.paragraphs, common_headers)

    txt_path = f"{filename}.txt"
    with open(os.path.join(base_path,txt_path), "w", encoding="utf-8") as txt_output:
        txt_output.write(remove_extra_newlines(paragraph_text))

    for i, table in enumerate(docx_input.tables, start=1):
        table_text = extract_text_from_table(table)
        if table_text:
            table_path = os.path.join(base_path,table_folder_path, f"{filename}_table{i}.txt")
            with open(table_path, "w", encoding="utf-8") as table_output:
                table_output.write(table_text)

    print(f"正文内容已保存到 {txt_path}")
    print(f"表格内容已保存到 {table_folder_path} 文件夹")

